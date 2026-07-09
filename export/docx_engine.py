# export/docx_engine.py
"""Génération DOCX professionnelle."""
from io import BytesIO
from loguru import logger
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from core.models import ExamPackage


class DOCXEngine:

    @classmethod
    def generate_docx(cls, exam: ExamPackage, is_answer_key: bool = False) -> bytes:
        doc = Document()

        # Configuration RTL
        style = doc.styles['Normal']
        style.font.name = 'Amiri'
        style.font.size = Pt(14)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Activer RTL
        for section in doc.sections:
            section.is_rtl = True

        # Titre
        title = doc.add_heading(exam.metadata.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Mode
        mode = "مفتاح التصحيح" if is_answer_key else "ورقة الاختبار"
        p = doc.add_paragraph(mode)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.color.rgb = RGBColor(0x1a, 0x73, 0xe8)

        # Métadonnées
        meta = exam.metadata
        doc.add_paragraph(f"المؤسسة: {meta.institution} | الأستاذ: {meta.teacher} | المدة: {meta.duration_minutes} دقيقة | المجموع: {exam.total_points} نقطة")

        # Champs étudiant
        if not is_answer_key:
            doc.add_paragraph("الاسم: ............................................    القسم: ....................    التاريخ: ....................")

        # Instructions
        doc.add_paragraph(f"📋 {meta.instructions}")

        # Texte
        doc.add_heading("📖 النص المشكّل", level=1)
        doc.add_paragraph(exam.vocalized_text)

        # QCM
        if exam.questions_mcq:
            doc.add_heading(f"I. أسئلة الفهم والمعجم ({sum(q.points for q in exam.questions_mcq)} نقطة)", level=1)
            for i, q in enumerate(exam.questions_mcq, 1):
                doc.add_paragraph(f"{i}. {q.question_text}  ({q.points} ن)")
                if not is_answer_key:
                    for j, opt in enumerate(q.options):
                        letter = "أبجد"[j] if j < 4 else str(j)
                        doc.add_paragraph(f"   {letter}) {opt.text}")
                else:
                    for opt in q.options:
                        if opt.is_correct:
                            p = doc.add_paragraph(f"   ✅ {opt.text}")
                            p.runs[0].font.color.rgb = RGBColor(0xc6, 0x28, 0x28)
                            p.runs[0].bold = True
                        else:
                            doc.add_paragraph(f"   {opt.text}")
                    if q.explanation:
                        doc.add_paragraph(f"   💡 {q.explanation}")

        # Cloze
        if exam.questions_cloze:
            doc.add_heading("II. الإعراب والصرف", level=1)
            for i, q in enumerate(exam.questions_cloze, 1):
                doc.add_paragraph(f"{i}. {q.cloze_sentence}  ({q.points} ن)")
                if is_answer_key:
                    doc.add_paragraph(f"   ✅ الإجابة: {q.answer_vocalized}")
                    doc.add_paragraph(f"   📐 القاعدة: {q.i3rab_rule}")
                    doc.add_paragraph(f"   📝 الإعراب: {q.i3rab_detailed}")
                else:
                    doc.add_paragraph("   الإجابة: ............................................")

        # Imlae
        if exam.questions_imlae:
            doc.add_heading("III. التصحيح الإملائي", level=1)
            for i, q in enumerate(exam.questions_imlae, 1):
                doc.add_paragraph(f"{i}. {q.sentence_with_errors}  ({q.points} ن)")
                if is_answer_key:
                    doc.add_paragraph(f"   ✅ التصحيح: {q.corrected_sentence}")
                    for err in q.errors:
                        doc.add_paragraph(f"   ❌ «{err.error_word}» → ✅ «{err.original_word}» — {err.rule_explanation}")

        # Vocabulaire
        if exam.questions_vocab:
            doc.add_heading("IV. المعجم", level=1)
            for i, v in enumerate(exam.questions_vocab, 1):
                doc.add_paragraph(f"{i}. اشرح كلمة «{v.vocalized}» في سياقها: {v.context_sentence}  ({v.points} ن)")
                if not is_answer_key:
                    doc.add_paragraph("   المعنى: ............................................")

        # Sauvegarde
        buf = BytesIO()
        doc.save(buf)
        logger.info("DOCX generated (answer_key={}) — {} bytes", is_answer_key, len(buf.getvalue()))
        return buf.getvalue()